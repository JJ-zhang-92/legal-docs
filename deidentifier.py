"""
Local PII de-identification engine for Chinese legal documents.
Zero external dependencies — pure regex + pattern matching.
Processes in milliseconds, designed for air-gapped execution.

Strategy: ROLE-AWARE replacement — identify entity's role in the document
first, then assign a consistent placeholder. Same entity always gets the
same tag across the entire document.

Usage:
    from deidentifier import deidentify
    result = deidentify(raw_text)
    print(result['deidentified_text'])  # safe to upload
    print(result['summary'])            # what was replaced
"""

import re
import json
import os
from collections import OrderedDict

# ═══════════════════════════════════════════════════════════════
# PATTERN DEFINITIONS
# ═══════════════════════════════════════════════════════════════

# Priority order: match longest/most specific first
PATTERNS = [
    # ── ID numbers (highest priority, most sensitive) ──
    ('身份证号', r'\b\d{6}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b'),
    ('统一社会信用代码', r'\b[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}\b'),
    ('工商注册号', r'\b\d{15}\b'),
    ('银行账号', r'\b\d{16,19}\b'),

    # ── Contact info ──
    ('手机号', r'\b1[3-9]\d{9}\b'),
    ('固定电话', r'\b0\d{2,3}[- ]?\d{7,8}(?:[- ]?\d{1,6})?\b'),
    ('邮箱', r'\b[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}\b'),

    # ── Case/legal identifier ──
    ('案号', r'[（(]\d{4}[）)]\S{2,15}号'),

    # ── Amounts (preserve magnitude) ──
    ('人民币金额', r'人民币\s*[\d,]+\.?\d*\s*[万亿千百]?\s*元'),
    ('美元金额', r'[$＄][\d,]+\.?\d*\s*[万亿千百]?\s*元?'),
    ('数字金额', r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\s*万元\b'),

    # ── Dates ──
    ('完整日期', r'\b\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日\b'),

    # ── Addresses ──
    # Single greedy pattern — avoids self-overlapping alternation fragments
    ('详细地址', r'[\u4e00-\u9fa5]{2,8}(?:省|自治区|特别行政区)?[\u4e00-\u9fa5a-zA-Z]{2,8}(?:市|自治州|地区)?[\u4e00-\u9fa5a-zA-Z]{2,8}(?:区|县|县级市|街道|镇)[\u4e00-\u9fa5a-zA-Z\d\-\s]{2,30}(?:号|号楼|楼|单元|室|层)?'),
]

# ═══════════════════════════════════════════════════════════════
# ROLE-AWARE ENTITY DETECTION
# ═══════════════════════════════════════════════════════════════

def detect_parties(text):
    """Detect party roles and their entity names."""
    parties = OrderedDict()

    role_patterns = [
        # (role_key, role_label, detection_regex)
        ('甲', '甲方', r'(?:甲方|供方|卖方|转让方|出租方|委托方|发包方|贷款人)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('乙', '乙方', r'(?:乙方|需方|买方|受让方|承租方|受托方|承包方|借款人)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('丙', '丙方', r'(?:丙方)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('原告', '原告', r'(?:原告)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('被告', '被告', r'(?:被告)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('上诉人', '上诉人', r'(?:上诉人)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('被上诉人', '被上诉人', r'(?:被上诉人)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
        ('申请人', '申请人', r'(?:申请人)[：:]\s*([\u4e00-\u9fa5（）()\w]+?)(?:[\n。；，]|$)'),
    ]

    for key, label, pattern in role_patterns:
        # Use finditer to capture all parties of the same role
        # (e.g., 原告一：张三，原告二：李四)
        seen_names = set()
        for m in re.finditer(pattern, text):
            name = m.group(1).strip()[:30]
            if name and name not in seen_names and len(name) >= 2:
                seen_names.add(name)
                suffix = f'_{len(seen_names)}' if len(seen_names) > 1 else ''
                parties[f'{key}{suffix}'] = {'name': name, 'label': label}

    # Detect company-like names not caught above
    # Use MULTILINE so ^ matches after every \n (not just start of string)
    # Use lookahead for trailing boundary to avoid consuming the next match
    company_pattern = re.findall(
        r'(?:^|\n|。|；|，)\s*([\u4e00-\u9fa5（）()]{2,20}?(?:有限公司|有限责任公司|股份有限公司|合伙企业|事务所|商行|厂|店))\s*(?=[，。\n]|$)',
        text, re.MULTILINE
    )

    for c in company_pattern:
        c = c.strip()
        already_known = any(c == v['name'] or v['name'] in c for v in parties.values())
        if not already_known and len(parties) < 6:
            parties[f'company_{len(parties)}'] = {'name': c, 'label': '其他公司'}

    return parties


# ═══════════════════════════════════════════════════════════════
# MAIN DE-IDENTIFICATION FUNCTION
# ═══════════════════════════════════════════════════════════════

def deidentify(text):
    """
    De-identify Chinese legal document text.

    Args:
        text: Raw document text (string)

    Returns:
        dict with keys:
            'deidentified_text': De-identified version (safe for AI to read)
            'original_length': Character count of original
            'deidentified_length': Character count of result
            'redaction_count': Number of entity types redacted
            'summary': Human-readable summary of what was done
            'entity_map': {placeholder: original_value} for local reference
            'parties': Detected party roles and names
            'paragraph_map': [{index, start, end, deid_text, original_text}, ...]
                            start/end are positions in ORIGINAL text.
                            Used by output script: AI reads deid_text → decides
                            which paragraphs → bash reads original[start:end] → docx.
    """
    result_text = text
    redactions = []
    entity_map = OrderedDict()
    party_map = detect_parties(text)

    # ── Step 0: Collect ALL matches from original text first ──
    all_matches = []  # (start, end, placeholder)

    # Party names
    for key, info in party_map.items():
        name = info['name']
        label = info['label']
        if len(name) >= 2:
            for m in re.finditer(re.escape(name), result_text):
                all_matches.append((m.start(), m.end(), f'[{label}]'))

    # Regex patterns
    for entity_type, pattern in PATTERNS:
        for m in re.finditer(pattern, result_text):
            all_matches.append((m.start(), m.end(), f'[{entity_type}已脱敏]'))

    # ── Step 0.5: Deduplicate overlapping spans (keep longest) ──
    # Overlapping matches break right-to-left replacement because span
    # positions become stale after the first replacement shifts text length.
    all_matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))  # by start asc, length desc
    deduped = []
    last_end = -1
    for start, end, placeholder in all_matches:
        if start >= last_end:
            deduped.append((start, end, placeholder))
            last_end = end
        # else: span overlaps with a longer span already selected → discard
    all_matches = deduped

    # ── Step 1: Apply replacements right-to-left (no overlaps remain) ──
    all_matches.sort(key=lambda x: x[0], reverse=True)

    counts = {}
    for start, end, placeholder in all_matches:
        count_key = placeholder
        counts[count_key] = counts.get(count_key, 0) + 1
        result_text = result_text[:start] + placeholder + result_text[end:]

    # ── Step 2: Build redaction summary ──
    for key, info in party_map.items():
        label = info['label']
        placeholder = f'[{label}]'
        cnt = counts.get(placeholder, 0)
        if cnt > 0:
            redactions.append(f'{label}名称 → {placeholder} ({cnt}处)')
            entity_map[placeholder] = info['name']

    for entity_type, pattern in PATTERNS:
        placeholder = f'[{entity_type}已脱敏]'
        cnt = counts.get(placeholder, 0)
        if cnt > 0:
            redactions.append(f'{entity_type} → {placeholder} ({cnt}处)')

    # ── Step 3: Build summary ──
    summary_lines = [
        f'原始文本: {len(text)} 字符',
        f'脱敏后: {len(result_text)} 字符',
        f'脱敏项数: {len(redactions)} 类',
        '---',
        '脱敏详情:',
    ]
    for r in redactions:
        summary_lines.append(f'  • {r}')
    if party_map:
        summary_lines.append('---')
        summary_lines.append('识别当事方:')
        for key, info in party_map.items():
            summary_lines.append(f'  [{info["label"]}] {info["name"]}')

    return {
        'deidentified_text': result_text,
        'original_length': len(text),
        'deidentified_length': len(result_text),
        'redaction_count': len(redactions),
        'summary': '\n'.join(summary_lines),
        'entity_map': entity_map,
        'parties': party_map,
        'paragraph_map': _build_paragraph_map(text, result_text),
    }


def _build_paragraph_map(original, deidentified):
    """Build paragraph-level position map: original text ↔ deidentified text.

    Each entry: {start, end, deid_text, original_text}
    - start/end: character positions in the ORIGINAL text
    - deid_text: the deidentified version of this paragraph
    - original_text: the ORIGINAL text of this paragraph

    This enables the bash output script to:
      AI reads deid_text → decides which paragraphs → bash reads
      original[pos][start:end] from disk → writes to docx
    """
    paras_orig = original.split('\n\n')
    paras_deid = deidentified.split('\n\n')

    # If split counts differ (e.g., deidentification merged some lines),
    # fall back to single newline split
    if len(paras_orig) != len(paras_deid):
        paras_orig = original.split('\n')
        paras_deid = deidentified.split('\n')

    # Guard against zip-length mismatch — if counts still differ after
    # fallback, use the shorter length to avoid silent truncation
    if len(paras_orig) != len(paras_deid):
        min_len = min(len(paras_orig), len(paras_deid))
        paras_orig = paras_orig[:min_len]
        paras_deid = paras_deid[:min_len]

    mapping = []
    pos = 0
    for i, (orig_p, deid_p) in enumerate(zip(paras_orig, paras_deid)):
        if orig_p.strip():
            start = original.index(orig_p, pos) if orig_p in original[pos:] else pos
            end = start + len(orig_p)
            pos = end
            mapping.append({
                'index': i,
                'start': start,
                'end': end,
                'deid_text': deid_p.strip(),
                'original_text': orig_p.strip(),
            })

    return mapping


def deidentify_file(filepath):
    """
    De-identify a .docx or .pdf file.

    Args:
        filepath: Path to document

    Returns:
        Same dict as deidentify(), plus 'source_file' key
    """
    import os
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.docx':
        from docx import Document
        doc = Document(filepath)
        text = '\n'.join(p.text for p in doc.paragraphs)
    elif ext == '.pdf':
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
    elif ext == '.txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
    else:
        raise ValueError(f'Unsupported file type: {ext}')

    result = deidentify(text)
    result['source_file'] = os.path.basename(filepath)
    return result


def deidentify_batch(filepaths):
    """
    De-identify multiple files in batch. Supports .docx, .pdf, .txt.

    Args:
        filepaths: List of file paths

    Returns:
        Dict mapping filename → deidentifier result dict
    """
    results = {}
    for fp in filepaths:
        fname = os.path.basename(fp)
        try:
            ext = os.path.splitext(fp)[1].lower()
            if ext == '.docx':
                from docx import Document
                doc = Document(fp)
                text = '\n'.join(p.text for p in doc.paragraphs)
            elif ext == '.pdf':
                import pdfplumber
                with pdfplumber.open(fp) as pdf:
                    text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
            elif ext == '.txt':
                with open(fp, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                results[fname] = {'error': f'Unsupported: {ext}', 'source_file': fname}
                continue

            deid = deidentify(text)
            deid['source_file'] = fname
            results[fname] = deid
        except Exception as e:
            results[fname] = {'error': str(e), 'source_file': fname}
    return results


if __name__ == '__main__':
    # Quick smoke test
    sample = """
甲方：浙江清研生物科技有限公司
统一社会信用代码：91330108MA28ABCDEF
法定代表人：李建国
联系电话：13812345678

乙方：义乌市清研日用品商行
经营者：王某某，女，1990年1月1日出生，汉族
身份证号：330782199001010001
住址：浙江省义乌市福田街道XX路123号

被告在拼多多平台销售侵权产品28653件，销售额约人民币1143254.70元。
原告于2024年3月15日进行公证证据保全。
案号：（2024）浙0108民初12345号
"""
    r = deidentify(sample)
    print(r['summary'].encode('utf-8', errors='replace').decode('utf-8'))
    print('\nDeidentified text (safe to upload):')
    print(r['deidentified_text'].encode('utf-8', errors='replace').decode('utf-8'))
