"""
Legal document template generators.
Produces properly formatted .docx files for Chinese legal documents.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


def _set_cell_border(cell, **kwargs):
    """Set cell border styles."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for attr in ('val', 'sz', 'color', 'space'):
                if attr in edge_data:
                    element.set(qn(f'w:{attr}'), str(edge_data[attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _safe_str(fields, key, default='________'):
    """Get field value, returning default if key is missing OR value is None."""
    v = fields.get(key)
    if v is None:
        return default
    return str(v)


def _safe_list(fields, key):
    """Get list field value, returning empty list if key is missing OR value is None.
    If the value is a string, wraps it in a list."""
    v = fields.get(key)
    if v is None:
        return []
    if isinstance(v, str):
        return [v]
    if isinstance(v, (list, tuple)):
        return list(v)
    return [str(v)]


def _add_heading_custom(doc, text, level=1):
    """Add a heading with Chinese-friendly formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    return heading


def _apply_document_style(doc):
    """Apply standard Chinese legal document formatting."""
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.paragraph_format.line_spacing = 1.5

    # Page margins: left 3.17cm for binding, others 2.54cm
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)


def _add_centered_line(doc, text, font_name='SimHei', font_size=Pt(16), bold=True):
    """Add a centered paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    return p


def _add_indented_paragraph(doc, text, first_line_indent=Cm(0.74)):
    """Add a body paragraph with 2-character first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)
    return p


def generate_complaint(output_path, **fields):
    """
    生成起诉状 (Civil Complaint).

    Required fields:
        plaintiff_name:   原告姓名/名称
        plaintiff_gender: 原告性别
        plaintiff_dob:    原告出生日期 (YYYY.MM.DD)
        plaintiff_ethnic: 原告民族
        plaintiff_address: 原告住址
        plaintiff_id:     原告身份证号
        plaintiff_phone:  原告联系电话
        defendant_name:   被告姓名/名称
        defendant_address: 被告住址
        defendant_phone:  被告联系电话
        court_name:       管辖法院名称
        claims:           诉讼请求 (list of strings)
        facts:            事实与理由 (string)
        evidence_list:    证据清单 (list of strings)
    """
    doc = Document()
    _apply_document_style(doc)

    # Title
    _add_centered_line(doc, '民 事 起 诉 状', 'SimHei', Pt(22), bold=True)

    # === 当事人信息 ===
    _add_heading_custom(doc, '当事人信息', level=2)

    # 原告
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('原告：')
    run.bold = True
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(14)

    _add_indented_paragraph(doc,
        f'姓名：{fields.get("plaintiff_name", "________")}　'
        f'性别：{fields.get("plaintiff_gender", "________")}　'
        f'出生日期：{fields.get("plaintiff_dob", "________")}'
    )
    _add_indented_paragraph(doc,
        f'民族：{fields.get("plaintiff_ethnic", "________")}　'
        f'身份证号：{fields.get("plaintiff_id", "________")}'
    )
    _add_indented_paragraph(doc,
        f'住址：{fields.get("plaintiff_address", "________")}'
    )
    _add_indented_paragraph(doc,
        f'联系电话：{fields.get("plaintiff_phone", "________")}'
    )

    # 被告
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('被告：')
    run.bold = True
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(14)

    if fields.get("defendant_is_company"):
        _add_indented_paragraph(doc,
            f'名称：{fields.get("defendant_name", "________")}'
        )
        _add_indented_paragraph(doc,
            f'法定代表人：{fields.get("defendant_legal_rep", "________")}'
        )
        _add_indented_paragraph(doc,
            f'住所地：{fields.get("defendant_address", "________")}'
        )
        _add_indented_paragraph(doc,
            f'统一社会信用代码：{fields.get("defendant_credit_code", "________")}'
        )
    else:
        _add_indented_paragraph(doc,
            f'姓名：{fields.get("defendant_name", "________")}　'
            f'性别：{fields.get("defendant_gender", "________")}'
        )
        _add_indented_paragraph(doc,
            f'出生日期：{fields.get("defendant_dob", "________")}　'
            f'民族：{fields.get("defendant_ethnic", "________")}'
        )
        _add_indented_paragraph(doc,
            f'住址：{fields.get("defendant_address", "________")}'
        )
        _add_indented_paragraph(doc,
            f'身份证号：{fields.get("defendant_id", "________")}'
        )
        _add_indented_paragraph(doc,
            f'联系电话：{fields.get("defendant_phone", "________")}'
        )

    # === 诉讼请求 ===
    _add_heading_custom(doc, '诉讼请求', level=2)
    claims = _safe_list(fields, "claims")
    for i, claim in enumerate(claims, 1):
        _add_indented_paragraph(doc, f'{i}. {claim}')

    # === 事实与理由 ===
    _add_heading_custom(doc, '事实与理由', level=2)
    facts = fields.get("facts", "（请在此陈述案件事实经过，包括时间、地点、人物、事件起因、经过、结果等要素。）")
    _add_indented_paragraph(doc, facts)

    # === 证据清单 ===
    _add_heading_custom(doc, '证据及证据来源', level=2)
    evidence = _safe_list(fields, "evidence_list")
    for i, ev in enumerate(evidence, 1):
        _add_indented_paragraph(doc, f'{i}. {ev}')

    # === 此致 ===
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('此 致')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(_safe_str(fields, "court_name", "________人民法院"))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    # Date and signature
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = fields.get("date", datetime.now().strftime('%Y 年 %m 月 %d 日'))
    run = p.add_run(f'具状人：{fields.get("plaintiff_name", "________")}')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(date_str)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    # 附项
    _add_heading_custom(doc, '附：', level=2)
    _add_indented_paragraph(doc, '1. 本起诉状副本____份；')
    _add_indented_paragraph(doc, '2. 证据材料____份共____页；')
    _add_indented_paragraph(doc, '3. 原告身份证复印件____份。')

    doc.save(output_path)
    return output_path


def generate_answer(output_path, **fields):
    """
    生成答辩状 (Answer/Reply Brief).

    Required fields:
        respondent_name:   答辩人姓名/名称
        respondent_address: 答辩人住址
        respondent_phone:  答辩人联系电话
        case_number:       案号
        court_name:        受诉法院名称
        claimant_name:     原告/上诉人名称
        response_points:   答辩意见列表 (list of strings)
        facts:             答辩事实与理由 (string)
    """
    doc = Document()
    _apply_document_style(doc)

    # Title
    _add_centered_line(doc, '民 事 答 辩 状', 'SimHei', Pt(22), bold=True)

    # 答辩人信息
    _add_heading_custom(doc, '答辩人信息', level=2)
    _add_indented_paragraph(doc,
        f'答辩人：{fields.get("respondent_name", "________")}'
    )
    _add_indented_paragraph(doc,
        f'住所地/住址：{fields.get("respondent_address", "________")}'
    )
    _add_indented_paragraph(doc,
        f'联系电话：{fields.get("respondent_phone", "________")}'
    )

    if fields.get("respondent_is_company"):
        _add_indented_paragraph(doc,
            f'法定代表人：{fields.get("respondent_legal_rep", "________")}'
        )
        _add_indented_paragraph(doc,
            f'统一社会信用代码：{fields.get("respondent_credit_code", "________")}'
        )

    # 案由
    _add_heading_custom(doc, '案由', level=2)
    _add_indented_paragraph(doc,
        f'答辩人就{fields.get("claimant_name", "________")}诉答辩人{fields.get("cause_of_action", "________")}一案'
        f'（案号：{fields.get("case_number", "________")}），提出答辩如下：'
    )

    # 答辩意见
    _add_heading_custom(doc, '答辩意见', level=2)
    response_points = _safe_list(fields, "response_points")
    for i, point in enumerate(response_points, 1):
        _add_indented_paragraph(doc, f'{i}. {point}')

    # 事实与理由
    _add_heading_custom(doc, '事实与理由', level=2)
    facts = fields.get("facts", "（请在此详述答辩的事实依据与法律理由。）")
    _add_indented_paragraph(doc, facts)

    # 此致
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('此 致')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(_safe_str(fields, "court_name", "________人民法院"))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    # 签名
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'答辩人：{fields.get("respondent_name", "________")}')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = fields.get("date", datetime.now().strftime('%Y 年 %m 月 %d 日'))
    run = p.add_run(date_str)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    _add_heading_custom(doc, '附：', level=2)
    _add_indented_paragraph(doc, '1. 本答辩状副本____份；')
    _add_indented_paragraph(doc, '2. 证据材料____份共____页。')

    doc.save(output_path)
    return output_path


def generate_legal_opinion(output_path, **fields):
    """
    生成法律意见书 (Legal Opinion Letter).

    Required fields:
        client_name:      委托人名称
        matter:           委托事项
        background:       案件基本事实
        legal_analysis:   法律分析 (string)
        risks:            风险提示 (list of strings)
        conclusion:       结论与建议 (string)
        law_firm:         律师事务所名称
        lawyer_name:      律师姓名
    """
    doc = Document()
    _apply_document_style(doc)

    # Title
    _add_centered_line(doc, '法 律 意 见 书', 'SimHei', Pt(22), bold=True)

    # 文书编号 (右上角)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fields.get("ref_number", f'〔{datetime.now().year}〕____字第____号'))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)

    # 委托人
    _add_heading_custom(doc, '一、基本情况', level=2)
    _add_indented_paragraph(doc, f'委托人：{fields.get("client_name", "________")}')
    _add_indented_paragraph(doc, f'委托事项：{fields.get("matter", "________")}')
    _add_indented_paragraph(doc, f'出具单位：{fields.get("law_firm", "________律师事务所")}')
    _add_indented_paragraph(doc, f'承办律师：{fields.get("lawyer_name", "________")}')

    # 事实认定
    _add_heading_custom(doc, '二、事实认定', level=2)
    _add_indented_paragraph(doc,
        fields.get("background", "（请在此陈述案件基本事实。）")
    )

    # 法律分析
    _add_heading_custom(doc, '三、法律分析', level=2)
    _add_indented_paragraph(doc,
        fields.get("legal_analysis", "（请在此进行详细的法律分析，包括适用的法律法规、司法解释、相关案例等。）")
    )

    # 风险提示
    _add_heading_custom(doc, '四、风险提示', level=2)
    risks = _safe_list(fields, "risks")
    for i, risk in enumerate(risks, 1):
        _add_indented_paragraph(doc, f'{i}. {risk}')

    # 结论与建议
    _add_heading_custom(doc, '五、结论与建议', level=2)
    _add_indented_paragraph(doc,
        fields.get("conclusion", "（请在此给出结论与具体建议。）")
    )

    # 声明
    _add_heading_custom(doc, '六、声明', level=2)
    _add_indented_paragraph(doc,
        '本法律意见书仅供委托人参考，未经本所书面同意，不得向第三方披露或用于其他目的。'
        '本意见基于现有材料出具，如发现新事实或证据，本所保留修改本意见的权利。'
    )

    # 落款
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fields.get("law_firm", "________律师事务所"))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'律师：{fields.get("lawyer_name", "________")}')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = fields.get("date", datetime.now().strftime('%Y 年 %m 月 %d 日'))
    run = p.add_run(date_str)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    doc.save(output_path)
    return output_path


def generate_demand_letter(output_path, **fields):
    """
    生成律师函 (Lawyer's Demand Letter).

    Required fields:
        sender_name:      发函方名称
        recipient_name:   收函方名称
        fact_statement:   事实陈述 (string)
        legal_basis:      法律依据 (string)
        demands:          主张/要求 (list of strings)
        deadline_days:    履行期限 (int, 天数)
        law_firm:         律师事务所名称
        lawyer_name:      律师姓名
    """
    doc = Document()
    _apply_document_style(doc)

    # Title
    _add_centered_line(doc, '律 师 函', 'SimHei', Pt(22), bold=True)

    # 编号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fields.get("ref_number", f'〔{datetime.now().year}〕____律函字第____号'))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)

    # 收函方
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'致：{fields.get("recipient_name", "________")}')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)
    run.bold = True

    # 委托声明
    _add_heading_custom(doc, '一、委托声明', level=2)
    _add_indented_paragraph(doc,
        f'{fields.get("law_firm", "________律师事务所")}（以下简称"本所"）接受'
        f'{fields.get("sender_name", "________")}（以下简称"委托人"）的委托，'
        f'指派{fields.get("lawyer_name", "________")}律师，就贵方{fields.get("matter", "________")}事宜，'
        f'特致函如下：'
    )

    # 事实陈述
    _add_heading_custom(doc, '二、基本事实', level=2)
    _add_indented_paragraph(doc,
        fields.get("fact_statement", "（请在此陈述相关事实。）")
    )

    # 法律依据
    _add_heading_custom(doc, '三、法律依据', level=2)
    _add_indented_paragraph(doc,
        fields.get("legal_basis", "（请在此列明适用的法律法规和合同条款。）")
    )

    # 律师意见与要求
    _add_heading_custom(doc, '四、律师意见与要求', level=2)
    demands = _safe_list(fields, "demands")
    for i, demand in enumerate(demands, 1):
        _add_indented_paragraph(doc, f'{i}. {demand}')

    deadline = int(fields.get("deadline_days") or 7)
    _add_indented_paragraph(doc,
        f'请贵方于收到本函之日起{deadline}日内履行上述义务。'
        f'如逾期未履行，委托人将保留通过诉讼、仲裁等法律途径维护自身合法权益的权利。'
    )

    # 落款
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(fields.get("law_firm", "________律师事务所"))
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'律师：{fields.get("lawyer_name", "________")}')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = fields.get("date", datetime.now().strftime('%Y 年 %m 月 %d 日'))
    run = p.add_run(date_str)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    # 联系方式
    _add_heading_custom(doc, '联系方式', level=2)
    _add_indented_paragraph(doc,
        f'地址：{fields.get("law_firm_address", "________")}'
    )
    _add_indented_paragraph(doc,
        f'电话：{fields.get("law_firm_phone", "________")}'
    )
    _add_indented_paragraph(doc,
        f'邮箱：{fields.get("law_firm_email", "________")}'
    )

    doc.save(output_path)
    return output_path


def generate_agency_agreement(output_path, **fields):
    """
    生成委托代理合同 (Agency/Retainer Agreement).

    Required fields:
        client_name:      委托人
        law_firm:         律师事务所
        lawyer_name:      承办律师
        matter:           委托事项
        fee_amount:       律师费金额
        fee_type:         收费方式 (固定/风险/计时)
    """
    doc = Document()
    _apply_document_style(doc)

    _add_centered_line(doc, '委 托 代 理 合 同', 'SimHei', Pt(20), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'（{datetime.now().year}）____律民代字第____号')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)

    # 合同双方
    _add_indented_paragraph(doc,
        f'甲方（委托人）：{fields.get("client_name", "________")}'
    )
    _add_indented_paragraph(doc,
        f'地址：{fields.get("client_address", "________")}'
    )
    _add_indented_paragraph(doc,
        f'电话：{fields.get("client_phone", "________")}'
    )

    doc.add_paragraph()
    _add_indented_paragraph(doc,
        f'乙方（受托人）：{fields.get("law_firm", "________律师事务所")}'
    )
    _add_indented_paragraph(doc,
        f'地址：{fields.get("law_firm_address", "________")}'
    )
    _add_indented_paragraph(doc,
        f'电话：{fields.get("law_firm_phone", "________")}'
    )
    _add_indented_paragraph(doc,
        f'承办律师：{fields.get("lawyer_name", "________")}'
    )

    # 鉴于条款
    _add_heading_custom(doc, '鉴于：', level=2)
    _add_indented_paragraph(doc,
        f'甲方因{fields.get("matter", "________")}事宜，委托乙方提供法律服务。'
        f'双方经协商一致，订立本合同，共同遵守。'
    )

    # 各条款
    sections = [
        ('第一条 委托事项', f'甲方委托乙方的法律事务为：{fields.get("matter", "________")}。'),
        ('第二条 承办律师', f'乙方指派{fields.get("lawyer_name", "________")}律师作为甲方的代理人，承办本合同约定的法律事务。'),
        ('第三条 律师费', fields.get("fee_clause",
            f'1. 律师费为人民币{fields.get("fee_amount", "________")}元。\n'
            f'2. 收费方式：{fields.get("fee_type", "固定收费")}。\n'
            f'3. 甲方应于本合同签订之日起____日内支付。')),
        ('第四条 办案费用', fields.get("expense_clause",
            '乙方律师办理甲方委托事务所发生的下列工作费用，由甲方承担：\n'
            '1. 行政、司法、鉴定、公证等部门收取的费用；\n'
            '2. 差旅费、食宿费；\n'
            '3. 翻译费、复印费、资料费等。')),
        ('第五条 甲方义务', fields.get("client_duties",
            '1. 真实、详尽、及时地向乙方律师叙述案情，提供相关证据材料；\n'
            '2. 积极配合乙方律师的工作；\n'
            '3. 按时足额支付律师费和工作费用。')),
        ('第六条 乙方义务', fields.get("firm_duties",
            '1. 勤勉尽责，维护甲方合法权益；\n'
            '2. 及时向甲方通报案件进展；\n'
            '3. 妥善保管甲方提供的证据材料；\n'
            '4. 对甲方提供的信息承担保密义务。')),
        ('第七条 合同解除', fields.get("termination_clause",
            '1. 经双方协商一致，可解除本合同；\n'
            '2. 甲方可以随时解除合同，但应按乙方已提供的服务收取费用；\n'
            '3. 乙方无正当理由不得解除合同。')),
        ('第八条 争议解决', fields.get("dispute_clause",
            '因本合同发生的争议，双方应协商解决；协商不成的，任何一方有权向乙方所在地人民法院提起诉讼。')),
    ]

    for title, content in sections:
        _add_heading_custom(doc, title, level=2)
        _add_indented_paragraph(doc, content)

    # 签署
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('甲方（签字/盖章）：                  乙方（盖章）：')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run('日期：                              日期：')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    doc.save(output_path)
    return output_path


def generate_power_of_attorney(output_path, **fields):
    """
    生成授权委托书 (Power of Attorney).

    Required fields:
        principal_name:   委托人
        attorney_name:    受托人/律师
        law_firm:         律师事务所
        case_matter:      案由
        opponent_name:    对方当事人
        scope:           代理权限 (一般代理 / 特别授权)
    """
    doc = Document()
    _apply_document_style(doc)

    _add_centered_line(doc, '授 权 委 托 书', 'SimHei', Pt(22), bold=True)

    _add_indented_paragraph(doc,
        f'委托人：{fields.get("principal_name", "________")}'
    )

    _add_indented_paragraph(doc,
        f'受托人：{fields.get("attorney_name", "________")}，'
        f'{fields.get("law_firm", "________律师事务所")}律师'
    )

    _add_indented_paragraph(doc,
        f'委托人因与{fields.get("opponent_name", "________")}'
        f'{fields.get("case_matter", "________")}一案，'
        f'现委托{fields.get("attorney_name", "________")}律师担任委托人的诉讼代理人。'
    )

    scope = fields.get("scope", "特别授权")
    if "特别" in scope:
        _add_indented_paragraph(doc,
            f'代理权限：特别授权代理。包括：代为承认、放弃、变更诉讼请求，进行和解，'
            f'提起反诉或上诉，代为签收法律文书等。'
        )
    else:
        _add_indented_paragraph(doc,
            f'代理权限：一般代理。包括：代为参与诉讼活动，提交证据材料，'
            f'发表代理意见，代为签收法律文书等。'
        )

    _add_indented_paragraph(doc,
        f'委托期限：{fields.get("term", "自签署之日起至本案终结之日止")}。'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'委托人（签字/盖章）：')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'受托人（签字）：')
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = fields.get("date", datetime.now().strftime('%Y 年 %m 月 %d 日'))
    run = p.add_run(date_str)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(14)

    doc.save(output_path)
    return output_path


TEMPLATE_REGISTRY = {
    '起诉状': {
        'generator': generate_complaint,
        'description': '民事起诉状 - 向法院提起诉讼的文书',
        'required': ['plaintiff_name', 'defendant_name', 'claims', 'facts', 'court_name'],
    },
    '答辩状': {
        'generator': generate_answer,
        'description': '民事答辩状 - 对起诉状进行答辩的文书',
        'required': ['respondent_name', 'case_number', 'response_points', 'facts', 'court_name'],
    },
    '法律意见书': {
        'generator': generate_legal_opinion,
        'description': '法律意见书 - 对法律问题出具专业意见',
        'required': ['client_name', 'matter', 'background', 'legal_analysis', 'conclusion'],
    },
    '律师函': {
        'generator': generate_demand_letter,
        'description': '律师函 - 向对方发出正式法律主张',
        'required': ['sender_name', 'recipient_name', 'fact_statement', 'demands'],
    },
    '委托代理合同': {
        'generator': generate_agency_agreement,
        'description': '委托代理合同 - 委托律师代理诉讼的合同',
        'required': ['client_name', 'law_firm', 'lawyer_name', 'matter', 'fee_amount'],
    },
    '授权委托书': {
        'generator': generate_power_of_attorney,
        'description': '授权委托书 - 向法院提交的代理授权文件',
        'required': ['principal_name', 'attorney_name', 'law_firm', 'case_matter', 'opponent_name'],
    },
}

if __name__ == '__main__':
    print("Available templates:")
    for name, info in TEMPLATE_REGISTRY.items():
        print(f"  {name}: {info['description']}")
        print(f"    Required fields: {', '.join(info['required'])}")
